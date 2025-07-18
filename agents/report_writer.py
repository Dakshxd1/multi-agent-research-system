from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import datetime
import os
import random
import re
import google.generativeai as genai

# Configure Gemini
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
llm = genai.GenerativeModel("gemini-2.0-pro")

def remove_bold_markdown(text):
    """Removes **bold markdown** without removing content."""
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)

import requests
from PIL import Image
from io import BytesIO

def download_wikipedia_image(topic, save_path="wiki_image.jpg"):
    """
    Downloads the first image from the Wikipedia page for the given topic.
    Returns path if successful, else None.
    """
    from bs4 import BeautifulSoup

    try:
        url = f"https://en.wikipedia.org/wiki/{topic.replace(' ', '_')}"
        response = requests.get(url)
        if response.status_code != 200:
            return None

        soup = BeautifulSoup(response.text, 'html.parser')
        infobox = soup.find("table", {"class": "infobox"})
        if not infobox:
            return None

        img = infobox.find("img")
        if not img:
            return None

        img_url = "https:" + img['src']
        img_data = requests.get(img_url).content
        with open(save_path, "wb") as f:
            f.write(img_data)

        return save_path
    except Exception as e:
        print(f"⚠️ Failed to download Wikipedia image: {e}")
        return None



def add_hyperlink(paragraph, url, text=None):
    """
    Adds a clickable hyperlink to a Word doc paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    text_run = OxmlElement("w:t")
    text_run.text = text or url
    new_run.append(text_run)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def generate_ai_conclusion(topic, verified_claims):
    try:
        summary_points = "\n".join([f"- {remove_bold_markdown(c['claim'])}" for c in verified_claims[:5]])
        prompt = f"""Write a professional conclusion summarizing a research report on "{topic}".
Include key insights based on the following claims:\n{summary_points}"""
        resp = llm.generate_content(prompt)
        return resp.text.strip()
    except Exception as e:
        print(f"⚠️ Gemini failed to generate conclusion: {e}")
        return f"This report explores various aspects of {topic} and shows how AI can summarize complex topics into concise insights."

def generate_report_docx(topic, intro, verified_claims, competitor_table="", filename="AI_Research_Report.docx"):
    doc = Document()
    doc.add_heading('AI Research Report', 0)
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_page_break()

    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(intro)
        # 👉 Insert Wikipedia image if available
    image_path = download_wikipedia_image(topic)
    if image_path:
        try:
            from docx.shared import Inches  # Make sure this is already at the top
            doc.add_picture(image_path, width=Inches(3.5))

            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # Center alignment
        except Exception as e:
            print(f"⚠️ Failed to insert image: {e}")
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # Center alignment
        except Exception as e:
            print(f"⚠️ Failed to insert image: {e}")


    doc.add_heading('Discussion / Main Sections', level=1)

    reword_prefixes = [
        "According to the article,",
        "The author emphasizes that",
        "The central idea conveyed is",
        "This source asserts that",
        "It is argued that",
        "As highlighted in the article,",
        "The article discusses how",
        "The key point presented is",
        "This report outlines that",
        "One significant insight is that",
        "The publication illustrates how",
        "Research suggests that",
        "It is emphasized that",
        "The content underscores the fact that",
        "The piece makes the case that"
    ]

    # Group by subtopic
    grouped = {}
    for claim in verified_claims:
        grouped.setdefault(claim.get('subtopic', 'General'), []).append(claim)

    for subtopic, claims in grouped.items():
        # Subtopic heading
        heading = doc.add_heading(level=2)
        run = heading.add_run(remove_bold_markdown(subtopic))
        run.font.size = Pt(14)

        for c in claims:
            cleaned = remove_bold_markdown(c['claim']).strip()

            if cleaned.lower().startswith("the main claim of the article is that"):
                trimmed = cleaned[len("the main claim of the article is that"):].strip()
                prefix = random.choice(reword_prefixes)
                final = f"{prefix} {trimmed[0].lower() + trimmed[1:]}"
            else:
                final = cleaned

            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(final)
            run.font.size = Pt(11)

            # Add hyperlink or fallback source
            source_url = c.get("source_url", "").strip()
            if source_url.lower().startswith("http"):
                p = doc.add_paragraph()
                add_hyperlink(p, source_url, "Source Link")
            else:
                fallback = doc.add_paragraph(f"Source: {source_url}", style='Intense Quote')
                for r in fallback.runs:
                    r.font.size = Pt(10)

    # Competitor Section
    doc.add_heading("Competitor Analysis", level=1)
    doc.add_paragraph(competitor_table or "No competitor data found.")

    # Conclusion Section
    doc.add_heading("Conclusion", level=1)
    ai_conclusion = generate_ai_conclusion(topic, verified_claims)
    doc.add_paragraph(ai_conclusion)

    # Save the file
    doc.save(filename)
    print(f"✅ Report saved as {filename}")
    return filename
